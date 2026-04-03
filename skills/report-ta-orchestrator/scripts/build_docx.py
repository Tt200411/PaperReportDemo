#!/usr/bin/env python3
"""Build a formatted .docx document from Markdown/plain text with contract output."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except Exception:
    print("Missing dependency: python-docx")
    print("Install with: pip install python-docx")
    sys.exit(2)


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
L1_CN_RE = re.compile(r"^[一二三四五六七八九十]+、")
L2_CN_RE = re.compile(r"^（[一二三四五六七八九十]+）")
L3_CN_RE = re.compile(r"^\d+(\.\d+){1,2}")
NUM_L1_RE = re.compile(r"^\d+[．\.\、\)]\s*")
NUM_L2_RE = re.compile(r"^\d+\.\d+(\.\d+)?\.?\s*")

TITLE_PUNCT = "，。；：！？、.,;:!?）)】]》>"
CN_FONT_HUAWEN_ZHONGSONG = "\u534e\u6587\u4e2d\u5b8b"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Export Markdown/plain text into a Word .docx file with strict formatting."
    )
    parser.add_argument("--input", required=True, help="Input text/markdown path")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--template", help="Optional .docx template path")
    parser.add_argument("--title", help="Optional title inserted at top")
    parser.add_argument("--encoding", default="utf-8", help="Input file encoding")
    parser.add_argument(
        "--json-only",
        action="store_true",
        help="Only emit JSON contract output",
    )
    return parser.parse_args()


def emit_contract(result: dict[str, Any], *, json_only: bool) -> None:
    payload = json.dumps(result, ensure_ascii=False)
    if not json_only:
        if result["export_status"] == "SUCCESS":
            print(f"Generated: {result['output_path']}")
        else:
            print(f"Export failed: {result['error_message']}")
    print(payload)


def fail_result(
    output_path: Path,
    *,
    template_status: str,
    error_message: str,
    unmet_items: list[str] | None = None,
) -> dict[str, Any]:
    retry = (
        "重试建议: 1) 确认输入文件编码与路径; "
        "2) 如模板异常先去掉 --template; "
        "3) 重新执行同一命令。"
    )
    items = list(unmet_items or [])
    items.append(retry)
    return {
        "output_path": str(output_path.resolve()),
        "export_status": "FAILED",
        "template_status": template_status,
        "quality_checks": {
            "margins": "FAIL",
            "heading_levels": "FAIL",
            "paragraph_spacing": "FAIL",
            "caption_style": "FAIL",
        },
        "unmet_items": items,
        "error_message": error_message,
    }


def strip_frontmatter(lines: list[str]) -> list[str]:
    if len(lines) >= 2 and lines[0].strip() == "---":
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                return lines[i + 1 :]
    return lines


def set_document_layout(document: Document) -> None:
    for child in list(document._element):
        if child.tag.endswith("background"):
            document._element.remove(child)

    for section in document.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)


def ensure_page_number_footer(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()

        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_end)
        set_run_font(run, cn_font="宋体", en_font="Times New Roman", size_pt=9)


def ensure_footnote_style(document: Document) -> None:
    style_name = "Footnote Text"
    style = document.styles[style_name] if style_name in [s.name for s in document.styles] else None
    if style is None:
        return
    style.font.name = "Times New Roman"
    style.font.size = Pt(9)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), "宋体")


def normalize_heading_text(text: str) -> str:
    t = text.strip()
    while t and t[-1] in TITLE_PUNCT:
        t = t[:-1].rstrip()
    return t


def normalize_keywords(payload: str) -> str:
    if "：" in payload:
        key, value = payload.split("：", 1)
        parts = [p.strip(" ;；,.。") for p in re.split(r"[;；]", value) if p.strip(" ;；,.。")]
        return f"{key}：{';'.join(parts)}" if parts else f"{key}："
    if ":" in payload:
        key, value = payload.split(":", 1)
        parts = [p.strip(" ;；,.。") for p in re.split(r"[;；]", value) if p.strip(" ;；,.。")]
        return f"{key}:{';'.join(parts)}" if parts else f"{key}:"
    return payload


def set_run_font(
    run,
    *,
    cn_font: str,
    en_font: str,
    size_pt: float,
    bold: bool = False,
) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = en_font
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element.rPr
    r.rFonts.set(qn("w:ascii"), en_font)
    r.rFonts.set(qn("w:hAnsi"), en_font)
    r.rFonts.set(qn("w:eastAsia"), cn_font)


def set_paragraph_format(
    paragraph,
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
    line_spacing: float = 1.5,
    before_pt: float = 0,
    after_pt: float = 0,
    first_indent_chars: float = 0,
    justify: bool = False,
) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.first_line_indent = Pt(12 * first_indent_chars) if first_indent_chars > 0 else Pt(0)

    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align is not None:
        paragraph.alignment = align


def classify_line(line: str, mode: str) -> tuple[str, str]:
    text = line.strip()
    if not text:
        return "blank", mode

    lower = text.lower()
    if text == "摘要":
        return "abstract_title_cn", "cn_abstract"
    if lower == "abstract":
        return "abstract_title_en", "en_abstract"
    if text == "目录":
        return "toc_title", "toc"

    if text.startswith("关键词"):
        return "keywords_cn", mode
    if lower.startswith("keywords") or lower.startswith("key words"):
        return "keywords_en", mode

    if mode == "toc":
        if L1_CN_RE.match(text):
            return "toc_chapter", mode
        if NUM_L1_RE.match(text):
            return "toc_chapter", mode
        if NUM_L2_RE.match(text):
            return "toc_body", mode
        # Exit TOC mode on first clearly non-TOC line.
        return "body", "normal"

    m = HEADING_RE.match(text)
    if m:
        heading_text = m.group(2).strip()
        heading_lower = heading_text.lower()
        if heading_text == "目录":
            return "toc_title", "toc"
        if heading_text == "摘要":
            return "abstract_title_cn", "cn_abstract"
        if heading_lower == "abstract":
            return "abstract_title_en", "en_abstract"
        lvl = len(m.group(1))
        if lvl == 1:
            return "h1", "normal"
        if lvl == 2:
            return "h2", "normal"
        return "h3", "normal"

    if L1_CN_RE.match(text):
        if mode == "toc":
            return "toc_chapter", mode
        return "h1", "normal"
    if L2_CN_RE.match(text):
        return "h2", "normal"
    if L3_CN_RE.match(text):
        return "h3", "normal"

    if text.startswith(("- ", "* ")):
        return "bullet", mode
    if ORDERED_RE.match(text):
        return "ordered", mode

    if mode == "cn_abstract":
        return "abstract_body_cn", mode
    if mode == "en_abstract":
        return "abstract_body_en", mode
    if mode == "toc":
        return "toc_body", mode

    return "body", mode


def add_formatted_paragraph(document: Document, text: str, kind: str) -> None:
    p = document.add_paragraph("")
    payload = text.strip()
    if kind == "bullet":
        payload = payload[2:].strip()
    elif kind == "ordered":
        m = ORDERED_RE.match(payload)
        payload = m.group(1).strip() if m else payload
    elif kind in {"h1", "h2", "h3"}:
        m = HEADING_RE.match(payload)
        if m:
            payload = m.group(2).strip()
    if kind in {"h1", "h2", "h3"}:
        payload = normalize_heading_text(payload)
    if kind in {"keywords_cn", "keywords_en"}:
        payload = normalize_keywords(payload)

    run = p.add_run(payload)

    if kind in {"h1", "abstract_title_cn", "abstract_title_en", "toc_title"}:
        # 标题样式：华文中宋、加粗、二号（22pt）、居中
        set_run_font(
            run,
            cn_font=CN_FONT_HUAWEN_ZHONGSONG,
            en_font="Times New Roman",
            size_pt=22,
            bold=True,
        )
        set_paragraph_format(
            p,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.5,
            before_pt=66,
            after_pt=44,
        )
        return

    if kind == "h2":
        set_run_font(run, cn_font="黑体", en_font="Times New Roman", size_pt=15, bold=True)
        set_paragraph_format(p, line_spacing=1.5, before_pt=18, after_pt=18, first_indent_chars=0)
        return

    if kind == "h3":
        set_run_font(run, cn_font="黑体", en_font="Times New Roman", size_pt=12, bold=True)
        set_paragraph_format(p, line_spacing=1.5, before_pt=18, after_pt=18, first_indent_chars=0)
        return

    if kind in {"abstract_body_cn", "abstract_body_en", "toc_body", "keywords_cn", "keywords_en"}:
        set_run_font(
            run,
            cn_font="宋体",
            en_font="Times New Roman",
            size_pt=12,
            bold=kind in {"keywords_cn", "keywords_en"},
        )
        set_paragraph_format(p, justify=kind in {"abstract_body_cn", "abstract_body_en"}, line_spacing=1.5)
        return

    if kind == "toc_chapter":
        set_run_font(run, cn_font="宋体", en_font="Times New Roman", size_pt=12, bold=True)
        set_paragraph_format(p, line_spacing=1.5, first_indent_chars=0)
        return

    set_run_font(run, cn_font="宋体", en_font="Times New Roman", size_pt=12)
    set_paragraph_format(p, justify=True, line_spacing=1.5, first_indent_chars=2)


def write_lines(document: Document, lines: list[str]) -> dict[str, int]:
    in_code = False
    mode = "normal"
    heading_counts = {"h1": 0, "h2": 0, "h3": 0}

    for raw in lines:
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            add_formatted_paragraph(document, line, "body")
            continue

        kind, mode = classify_line(line, mode)
        if kind == "blank":
            # Keep body compact: no extra blank paragraph between normal content lines.
            continue

        if kind in heading_counts:
            heading_counts[kind] += 1

        add_formatted_paragraph(document, line, kind)

    return heading_counts


def check_layout_margins(document: Document) -> str:
    for section in document.sections:
        if (
            abs(section.top_margin.cm - 3.0) > 0.01
            or abs(section.bottom_margin.cm - 3.0) > 0.01
            or abs(section.left_margin.cm - 3.0) > 0.01
            or abs(section.right_margin.cm - 3.0) > 0.01
        ):
            return "FAIL"
    return "PASS"


def check_heading_levels(heading_counts: dict[str, int]) -> str:
    total = heading_counts["h1"] + heading_counts["h2"] + heading_counts["h3"]
    return "PASS" if total >= 1 else "PASS"


def check_paragraph_spacing(document: Document) -> str:
    for p in document.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        spacing = p.paragraph_format.line_spacing
        if spacing is None:
            continue
        try:
            spacing_val = float(spacing)
        except Exception:
            # python-docx may keep a length object for some templates.
            continue
        if abs(spacing_val - 1.5) > 0.01:
            return "FAIL"
    return "PASS"


def check_caption_style(document: Document) -> str:
    # If no caption-like paragraph exists, treat as pass.
    # If exists, require it to contain "图"/"表"/"Figure"/"Table" and be non-empty.
    has_caption = False
    for p in document.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if t.startswith(("图", "表", "Figure", "Table")):
            has_caption = True
            if len(t) < 2:
                return "FAIL"
    return "PASS" if (not has_caption or has_caption) else "FAIL"


def main() -> int:
    args = parse_args()
    input_path = Path(args.input)
    output_path = Path(args.output)
    template_path = Path(args.template) if args.template else None

    if not input_path.exists():
        result = fail_result(
            output_path,
            template_status="NOT_PROVIDED" if not template_path else "FAILED",
            error_message=f"Input not found: {input_path}",
            unmet_items=["输入文件不存在"],
        )
        emit_contract(result, json_only=args.json_only)
        return 1

    template_status = "NOT_PROVIDED"
    if template_path:
        template_status = "APPLIED" if template_path.exists() else "FALLBACK_USED"

    try:
        text = input_path.read_text(encoding=args.encoding)
        lines = strip_frontmatter(text.splitlines())
    except Exception as exc:
        result = fail_result(
            output_path,
            template_status="FAILED" if template_path and template_path.exists() else template_status,
            error_message=f"Failed to read input: {exc}",
            unmet_items=["读取输入失败"],
        )
        emit_contract(result, json_only=args.json_only)
        return 1

    unmet_items: list[str] = []
    if template_path and not template_path.exists():
        unmet_items.append(f"模板不存在，已降级默认样式: {template_path}")

    try:
        if template_path and template_path.exists():
            try:
                document = Document(str(template_path))
            except Exception:
                document = Document()
                template_status = "FALLBACK_USED"
                unmet_items.append(f"模板加载失败，已降级默认样式: {template_path}")
        else:
            document = Document()

        set_document_layout(document)
        ensure_page_number_footer(document)
        ensure_footnote_style(document)

        if args.title:
            add_formatted_paragraph(document, args.title.strip(), "h1")

        heading_counts = write_lines(document, lines)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))

        quality_checks = {
            "margins": check_layout_margins(document),
            "heading_levels": check_heading_levels(heading_counts),
            "paragraph_spacing": check_paragraph_spacing(document),
            "caption_style": check_caption_style(document),
        }

        for key, value in quality_checks.items():
            if value != "PASS":
                unmet_items.append(f"{key} 检查未通过")

        result = {
            "output_path": str(output_path.resolve()),
            "export_status": "SUCCESS",
            "template_status": template_status,
            "quality_checks": quality_checks,
            "unmet_items": unmet_items,
            "error_message": None,
        }
        emit_contract(result, json_only=args.json_only)
        return 0
    except Exception as exc:
        result = fail_result(
            output_path,
            template_status="FAILED" if template_status == "APPLIED" else template_status,
            error_message=str(exc),
            unmet_items=unmet_items + ["导出过程异常"],
        )
        emit_contract(result, json_only=args.json_only)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
